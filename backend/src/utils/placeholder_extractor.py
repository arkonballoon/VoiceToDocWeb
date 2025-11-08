"""
Utility-Modul für die Extraktion von Platzhaltern aus verschiedenen Dateiformaten
"""
import re
from typing import List, Set, Dict, Optional
from pathlib import Path
from utils.logger import get_logger

logger = get_logger(__name__)

class PlaceholderExtractor:
    """Extrahiert Platzhalter aus verschiedenen Dateiformaten"""
    
    # Regex-Pattern für Platzhalter im Format {{feldname}}
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
    
    @staticmethod
    def extract_from_text(text: str) -> Set[str]:
        """
        Extrahiert Platzhalter aus einem Text-String
        
        Args:
            text: Der Text, aus dem Platzhalter extrahiert werden sollen
            
        Returns:
            Set von Platzhalternamen (ohne {{}})
        """
        matches = PlaceholderExtractor.PLACEHOLDER_PATTERN.findall(text)
        # Entferne Leerzeichen und normalisiere
        placeholders = {match.strip() for match in matches if match.strip()}
        return placeholders
    
    @staticmethod
    def extract_from_docx(file_path: Path) -> Set[str]:
        """
        Extrahiert Platzhalter aus einer Word-Datei (.docx)
        
        Args:
            file_path: Pfad zur .docx-Datei
            
        Returns:
            Set von Platzhalternamen
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            all_text = []
            
            # Durchlaufe alle Absätze
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            # Durchlaufe alle Tabellen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            
            # Durchlaufe Header und Footer
            for section in doc.sections:
                for header in section.header.paragraphs:
                    all_text.append(header.text)
                for footer in section.footer.paragraphs:
                    all_text.append(footer.text)
            
            text = '\n'.join(all_text)
            return PlaceholderExtractor.extract_from_text(text)
            
        except ImportError:
            logger.error("python-docx ist nicht installiert")
            raise ImportError("python-docx ist erforderlich für .docx-Unterstützung")
        except Exception as e:
            logger.error(f"Fehler beim Lesen der Word-Datei: {str(e)}")
            raise
    
    @staticmethod
    def extract_from_xlsx(file_path: Path) -> Set[str]:
        """
        Extrahiert Platzhalter aus einer Excel-Datei (.xlsx)
        
        Args:
            file_path: Pfad zur .xlsx-Datei
            
        Returns:
            Set von Platzhalternamen
        """
        try:
            from openpyxl import load_workbook
            
            all_text = []
            
            # WICHTIG: Mit data_only=False liest openpyxl Formeln als Strings (z.B. "=A1+B1")
            # Mit data_only=True liest es nur die berechneten Werte
            # Wir brauchen BEIDES: Formeln (können Platzhalter enthalten) UND berechnete Werte
            try:
                workbook = load_workbook(file_path, data_only=False)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    logger.debug(f"Verarbeite Arbeitsblatt: {sheet_name} (Formeln und Werte)")
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value is not None:
                                cell_str = str(cell.value)
                                all_text.append(cell_str)
                                # Wenn es eine Formel ist (beginnt mit =), ist das der Formel-Text
                                if cell.data_type == 'f' or (isinstance(cell.value, str) and cell.value.startswith('=')):
                                    logger.debug(f"Formel gefunden: {cell_str[:200]}")
            except Exception as e:
                logger.warning(f"Fehler beim Lesen von Formeln: {str(e)}")
                logger.exception(e)
            
            # Dann mit data_only=True für berechnete Werte (falls vorhanden)
            try:
                workbook = load_workbook(file_path, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows(values_only=True):
                        for cell_value in row:
                            if cell_value is not None:
                                all_text.append(str(cell_value))
            except Exception as e:
                logger.warning(f"Fehler beim Lesen berechneter Werte: {str(e)}")
            
            # Debug: Zeige ersten 500 Zeichen des extrahierten Textes
            text = '\n'.join(all_text)
            logger.debug(f"Extrahierter Text aus XLSX (erste 500 Zeichen): {text[:500]}")
            logger.debug(f"Gesamte Textlänge: {len(text)} Zeichen")
            
            placeholders = PlaceholderExtractor.extract_from_text(text)
            logger.info(f"Extrahierte Platzhalter aus XLSX: {placeholders}")
            
            # Wenn keine Platzhalter gefunden wurden, zeige Beispiel-Zellen
            if not placeholders:
                logger.warning(f"Keine Platzhalter im Format {{{{feldname}}}} gefunden in {file_path}")
                # Zeige erste 10 nicht-leere Zellen als Hinweis
                sample_cells = [t for t in all_text[:20] if t and len(t.strip()) > 0]
                if sample_cells:
                    logger.info(f"Beispiel-Zellen-Inhalte (erste 10): {sample_cells[:10]}")
            
            return placeholders
            
        except ImportError:
            logger.error("openpyxl ist nicht installiert")
            raise ImportError("openpyxl ist erforderlich für .xlsx-Unterstützung")
        except Exception as e:
            logger.error(f"Fehler beim Lesen der Excel-Datei: {str(e)}")
            raise
    
    @staticmethod
    def extract_from_file(file_path: Path) -> Set[str]:
        """
        Extrahiert Platzhalter aus einer Datei basierend auf der Dateierweiterung
        
        Args:
            file_path: Pfad zur Datei
            
        Returns:
            Set von Platzhalternamen
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return PlaceholderExtractor.extract_from_docx(file_path)
        elif suffix == '.xlsx':
            return PlaceholderExtractor.extract_from_xlsx(file_path)
        elif suffix in ['.md', '.txt', '.markdown']:
            # Für Text-Dateien einfach den Inhalt lesen
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return PlaceholderExtractor.extract_from_text(text)
        else:
            raise ValueError(f"Nicht unterstütztes Dateiformat: {suffix}")
