import re
from typing import Dict, List, Optional, Any
from pydantic import BaseModel
from utils.logger import get_logger, log_function_call
import backoff
from openai import AsyncOpenAI
import logging
import json
from config import settings
from fastapi import WebSocket
import asyncio
from datetime import datetime
from utils.singleton import Singleton
from pathlib import Path
import uuid

logger = get_logger(__name__)

class TemplateProcessingResult(BaseModel):
    """Ergebnis der Template-Verarbeitung"""
    class Config:
        arbitrary_types_allowed = True
    
    extracted_info: Dict[str, str]
    filled_template: str
    validation_result: Optional[Dict[str, Any]] = None
    needs_revision: bool = False
    revision_comments: Optional[str] = None

class TemplateProcessor(Singleton):
    def _init(self):
        """Initialisierung des TemplateProcessors"""
        self.api_key = settings.LLM_API_KEY
        self.model = settings.LLM_MODEL
        self.client = AsyncOpenAI(api_key=self.api_key)
        self.active_connections = {}
        logger.info(f"TemplateProcessor initialisiert mit API-Key: {self.api_key[:10]}...")
        logger.info(f"Verwende LLM-Modell: {self.model}")
    
    async def register_connection(self, process_id: str, websocket: WebSocket):
        """Registriert eine neue WebSocket-Verbindung für Updates"""
        self.active_connections[process_id] = websocket
        
    async def remove_connection(self, process_id: str):
        """Entfernt eine WebSocket-Verbindung"""
        if process_id in self.active_connections:
            del self.active_connections[process_id]
    
    async def send_update(self, process_id: str, status: str, progress: float, message: str):
        """Sendet ein Update über WebSocket"""
        if process_id in self.active_connections:
            try:
                websocket = self.active_connections[process_id]
                update = {
                    "type": "template_update",
                    "status": status,
                    "progress": progress,
                    "message": message,
                    "timestamp": datetime.now().isoformat()
                }
                
                await websocket.send_json(update)
                logger.debug(f"Update gesendet: {update}")
                
            except Exception as e:
                logger.error(f"Fehler beim Senden des Updates: {str(e)}")
                await self.remove_connection(process_id)
    
    async def process_template_with_updates(
        self, 
        template: str, 
        transcription: str,
        process_id: str,
        placeholders: Optional[Dict[str, str]] = None,
        template_obj: Optional[Any] = None,
        output_dir: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Verarbeitet ein Template mit der gegebenen Transkription und sendet Updates"""
        try:
            logger.info(f"Starte Template-Verarbeitung für {process_id}")
            await self.send_update(
                process_id, 
                "started", 
                0.0, 
                "Starte Verarbeitung der Vorlage..."
            )
            
            # Informationen extrahieren
            logger.info(f"Extrahiere Informationen für {process_id}")
            await self.send_update(
                process_id,
                "extracting",
                0.25,
                "Extrahiere Informationen aus der Transkription..."
            )
            extracted_info = await self._extract_information(template, transcription, placeholders)
            
            # Template füllen
            logger.info(f"Fülle Template für {process_id}")
            await self.send_update(
                process_id,
                "filling",
                0.50,
                "Verarbeite Template mit extrahierten Informationen..."
            )
            filled_template = await self._fill_template(template, extracted_info, transcription)
            
            # Validierung
            logger.info(f"Validiere Ergebnis für {process_id}")
            await self.send_update(
                process_id,
                "validating",
                0.75,
                "Validiere das ausgefüllte Template..."
            )
            validation_result = await self._validate_result(
                template, transcription, extracted_info, filled_template
            )
            
            result = {
                "processed_text": filled_template,
                "extracted_info": extracted_info,
                "validation_result": json.loads(validation_result),
                "metadata": {
                    "model": settings.LLM_MODEL,
                    "timestamp": datetime.now().isoformat()
                }
            }
            
            # Format-spezifische Ausgabe erstellen (wenn Template-Objekt und output_dir vorhanden)
            if template_obj and output_dir:
                await self.send_update(
                    process_id,
                    "creating_output",
                    0.90,
                    "Erstelle format-spezifische Ausgabedatei..."
                )
                output_file_path = await self._create_formatted_output(
                    template_obj, extracted_info, process_id, output_dir
                )
                if output_file_path:
                    result["output_file_path"] = output_file_path
                    logger.info(f"Format-spezifische Ausgabe erstellt: {output_file_path}")
            
            await self.send_update(
                process_id, 
                "completed", 
                1.0, 
                "Verarbeitung erfolgreich abgeschlossen"
            )
            logger.info(f"Template-Verarbeitung für {process_id} abgeschlossen")
            
            return result
            
        except Exception as e:
            logger.error(f"Fehler bei Template-Verarbeitung {process_id}: {str(e)}")
            await self.send_update(
                process_id, 
                "error", 
                1.0, 
                f"Fehler bei der Verarbeitung: {str(e)}"
            )
            raise
        finally:
            # Verbindung entfernen
            await self.remove_connection(process_id)
    
    async def process_template(
        self, 
        template: str, 
        transcription: str,
        placeholders: Optional[Dict[str, str]] = None,
        template_obj: Optional[Any] = None,
        output_dir: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Verarbeitet ein Template mit der gegebenen Transkription"""
        try:
            # Parallele Verarbeitung mit Timeout
            tasks = [
                asyncio.create_task(self._extract_information(template, transcription, placeholders)),
                asyncio.create_task(self._fill_template(template, {}, transcription)),
            ]
            
            # Warte maximal 60 Sekunden auf alle Tasks
            results = await asyncio.wait_for(asyncio.gather(*tasks), timeout=60.0)
            extracted_info, filled_template = results

            # Validierung separat durchführen
            validation_result = await self._validate_result(
                template,
                transcription,
                extracted_info,
                filled_template
            )

            result = {
                "processed_text": filled_template,
                "extracted_info": extracted_info,
                "validation_result": json.loads(validation_result),
                "metadata": {
                    "model": settings.LLM_MODEL,
                    "timestamp": datetime.now().isoformat()
                }
            }
            
            # Format-spezifische Ausgabe erstellen (wenn Template-Objekt und output_dir vorhanden)
            if template_obj and output_dir:
                output_file_path = await self._create_formatted_output(
                    template_obj, extracted_info, str(uuid.uuid4()), output_dir
                )
                if output_file_path:
                    result["output_file_path"] = output_file_path
            
            return result

        except asyncio.TimeoutError:
            logger.error("Timeout bei der Template-Verarbeitung")
            raise TimeoutError("Die Verarbeitung hat zu lange gedauert")
            
        except Exception as e:
            logger.error(f"Fehler bei der Template-Verarbeitung: {str(e)}")
            raise
    
    async def _extract_information(
        self, 
        template: str, 
        transcription: str,
        placeholders: Optional[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        Extrahiert Informationen aus der Transkription basierend auf Platzhaltern
        
        Args:
            template: Template-Text
            transcription: Transkriptionstext
            placeholders: Dictionary mit Platzhalternamen -> Prompts
        """
        try:
            # Wenn Platzhalter vorhanden sind, verwende diese für die Extraktion
            if placeholders:
                # Erstelle eine strukturierte Prompt-Liste für jeden Platzhalter
                extraction_prompts = []
                for placeholder_name, prompt in placeholders.items():
                    extraction_prompts.append(f"- {placeholder_name}: {prompt}")
                
                system_prompt = f"""
                    Extrahiere die folgenden Informationen aus der Transkription basierend auf den gegebenen Prompts.
                    Gib nur die gefundenen Informationen als JSON zurück.
                    Format: {{"platzhaltername": "extrahierter_wert"}}
                    Wenn eine Information nicht gefunden werden kann, verwende einen leeren String "".
                    Bei Listen gebe die Werte als kommagetrennte Zeichenkette zurück.
                    
                    Zu extrahierende Informationen:
                    {chr(10).join(extraction_prompts)}
                """
            else:
                # Fallback auf alte Methode für Rückwärtskompatibilität
                system_prompt = """
                    Extrahiere die im Template-Header unter "## Benötigte Informationen" 
                    markierten Informationen aus der Transkription. 
                    Gib nur die gefundenen Informationen als JSON zurück.
                    Format: {"field_name": "extracted_value"}
                    Wenn eine Information nicht gefunden werden kann, verwende einen leeren String "".
                    Bei Listen gebe die Werte als kommagetrennte Zeichenkette zurück.
                """
            
            response = await asyncio.wait_for(
                self.client.chat.completions.create(
                    model=settings.LLM_MODEL,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Template:\n{template}\n\nTranskription:\n{transcription}"}
                    ],
                    response_format={ "type": "json_object" },
                    temperature=0.3
                ),
                timeout=30.0
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            logger.error(f"Fehler bei der Informationsextraktion: {str(e)}")
            return {}
    
    async def _fill_template(
        self, 
        template: str, 
        extracted_info: Dict[str, str],
        additional_context: Optional[str] = None
    ) -> str:
        """Füllt das Template mit den extrahierten Informationen"""
        context = f"Additional Context:\n{additional_context}\n\n" if additional_context else ""
        
        response = await self.client.chat.completions.create(
                model=settings.LLM_MODEL,
            messages=[
                {"role": "system", "content": """
                    Fülle das Template mit den extrahierten Informationen aus.
                    Verwende die gegebenen Informationen.
                    Verwende dabei die Struktur des Templates ab ##Struktur. 
                    Ersetze den Text "Beschreibung: ..." in den einzelnen Abschnitten durch einen vollständigen Text, 
                    der die Beschreibung des Abschnitts umsetzt und die extrahierten Informationen enthält.
                    Die Überschriften (###) und Struktur müssen beibehalten werden.
                    Formatiere den Text professionell und lesbar.
                """},
                {"role": "user", "content": f"""
                    Template:\n{template}\n\n
                    Extrahierte Informationen:\n{extracted_info}\n\n
                    {context}
                """}
            ]
        )
        content = response.choices[0].message.content
        # Entferne "Beschreibung: " am Zeilenanfang
        cleaned_content = re.sub(r'(?m)^Beschreibung:\s*', '', content)
        return cleaned_content
    
    async def _validate_result(
        self, 
        template: str, 
        transcription: str, 
        extracted_info: Dict[str, str],
        filled_template: str
    ) -> Dict[str, any]:
        """Validiert das ausgefüllte Template"""
        response = await self.client.chat.completions.create(
                model=settings.LLM_MODEL,
            messages=[
                {"role": "system", "content": """
                    Überprüfe das ausgefüllte Template auf:
                    1. Vollständigkeit der benötigten Informationen
                    2. Korrekte Verwendung der extrahierten Informationen
                    3. Einhaltung der Template-Struktur
                    4. Konsistenz mit der Original-Transkription
                    
                   Gib das Ergebnis als JSON zurück mit exakt dieser Struktur:
                    {
                        "is_valid": boolean,
                        "needs_revision": boolean,
                        "revision_comments": string,
                        "validation_details": {
                            "completeness_score": float,
                            "structure_score": float,
                            "consistency_score": float,
                            "missing_fields": string[],
                            "structure_issues": string[],
                            "consistency_issues": string[]
                        },
                        "improvement_suggestions": {
                            "general_feedback": string,
                            "specific_suggestions": string[]
                        }
                    }
                     Wichtig: 
                     - Alle numerischen Werte müssen zwischen 0.0 und 1.0 liegen
                     - Arrays können leer sein, aber müssen immer als Array existieren
                     - Gebe konkrete, actionable Verbesserungsvorschläge
                     - Das general_feedback sollte eine Zusammenfassung der wichtigsten Punkte sein
                     - specific_suggestions sollte spezifische, umsetzbare Vorschläge enthalten
                     - Antworte in Deutsch
                """},
                {"role": "user", "content": f"""
                    Original Template:\n{template}\n\n
                    Transkription:\n{transcription}\n\n
                    Extrahierte Informationen:\n{extracted_info}\n\n
                    Ausgefülltes Template:\n{filled_template}
                """}
            ],
            response_format={ "type": "json_object" }
        )
        return response.choices[0].message.content 
    
    async def _create_formatted_output(
        self,
        template_obj: Any,
        extracted_info: Dict[str, str],
        process_id: str,
        output_dir: Path
    ) -> Optional[str]:
        """
        Erstellt eine format-spezifische Ausgabedatei (Word/Excel) mit ausgefüllten Platzhaltern
        
        Args:
            template_obj: Template-Objekt mit file_format und file_path
            extracted_info: Extrahierte Informationen (Platzhalter -> Werte)
            process_id: Prozess-ID für Dateinamen
            output_dir: Verzeichnis für Ausgabedateien
            
        Returns:
            Pfad zur erstellten Datei oder None bei Fehler/Markdown
        """
        try:
            # Prüfe, ob format-spezifische Ausgabe benötigt wird
            file_format = getattr(template_obj, 'file_format', None)
            file_path = getattr(template_obj, 'file_path', None)
            
            logger.info(f"[PLATZHALTER-ERSETZUNG] Starte format-spezifische Ausgabe für Prozess {process_id}")
            logger.debug(f"[PLATZHALTER-ERSETZUNG] Dateiformat: {file_format}, Dateipfad: {file_path}")
            logger.debug(f"[PLATZHALTER-ERSETZUNG] Extrahierte Informationen: {extracted_info}")
            logger.info(f"[PLATZHALTER-ERSETZUNG] Anzahl extrahierter Werte: {len(extracted_info)}")
            
            if not file_format or not file_path:
                logger.debug("Kein file_format oder file_path vorhanden, überspringe format-spezifische Ausgabe")
                return None
            
            if file_format == 'markdown':
                logger.debug("Markdown-Template, überspringe format-spezifische Ausgabe")
                return None
            
            # Prüfe, ob Originaldatei existiert
            original_path = Path(file_path)
            if not original_path.exists():
                logger.warning(f"Originaldatei nicht gefunden: {file_path}, Fallback auf Text-Ausgabe")
                return None
            
            # Erstelle Ausgabedateinamen
            template_name = getattr(template_obj, 'name', 'template')
            # Bereinige Dateinamen (entferne ungültige Zeichen)
            safe_name = re.sub(r'[^\w\s-]', '', template_name).strip()
            safe_name = re.sub(r'[-\s]+', '_', safe_name)
            extension = original_path.suffix
            output_filename = f"{process_id}_{safe_name}{extension}"
            output_path = output_dir / output_filename
            
            logger.info(f"Erstelle format-spezifische Ausgabe: {file_format} -> {output_path}")
            
            if file_format == 'docx':
                return await self._create_docx_output(original_path, output_path, extracted_info, process_id)
            elif file_format == 'xlsx':
                return await self._create_xlsx_output(original_path, output_path, extracted_info, process_id)
            else:
                logger.warning(f"Unbekanntes Dateiformat: {file_format}")
                return None
                
        except Exception as e:
            logger.error(f"Fehler beim Erstellen der format-spezifischen Ausgabe: {str(e)}")
            logger.exception(e)
            return None
    
    async def _create_docx_output(
        self,
        original_path: Path,
        output_path: Path,
        extracted_info: Dict[str, str],
        process_id: str
    ) -> Optional[str]:
        """Erstellt eine Word-Datei mit ausgefüllten Platzhaltern"""
        try:
            from docx import Document
            
            logger.info(f"[WORD-ERSETZUNG] Lade Word-Datei: {original_path}")
            logger.debug(f"[WORD-ERSETZUNG] Verfügbare Werte: {extracted_info}")
            
            # Lade Originaldokument
            doc = Document(str(original_path))
            
            replacements_count = 0
            
            # Ersetze Platzhalter in Absätzen
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text:
                    original_text = paragraph.text
                    # Finde alle Platzhalter im Text
                    placeholder_pattern = r'\{\{([^}]+)\}\}'
                    found_placeholders_in_text = re.findall(placeholder_pattern, original_text)
                    if found_placeholders_in_text:
                        logger.debug(f"[WORD-ERSETZUNG] Absatz {para_idx}: Gefundene Platzhalter: {found_placeholders_in_text}")
                        logger.debug(f"[WORD-ERSETZUNG] Absatz {para_idx}: Vollständiger Text: {original_text[:500]}")
                    if self._has_placeholders(original_text):
                        logger.debug(f"[WORD-ERSETZUNG] Absatz {para_idx}: Gefundener Text mit Platzhaltern: {original_text[:200]}")
                        new_text = self._replace_placeholders(original_text, extracted_info, process_id)
                        if new_text != original_text:
                            logger.info(f"[WORD-ERSETZUNG] Absatz {para_idx}: Ersetzt '{original_text[:100]}' -> '{new_text[:100]}'")
                            replacements_count += 1
                            # Ersetze Text, behalte Formatierung
                            for run in paragraph.runs:
                                original_run_text = run.text
                                run.text = self._replace_placeholders(run.text, extracted_info, process_id)
                                if run.text != original_run_text:
                                    logger.debug(f"[WORD-ERSETZUNG] Run ersetzt: '{original_run_text[:50]}' -> '{run.text[:50]}'")
            
            # Ersetze Platzhalter in Tabellen
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text:
                            original_cell_text = cell.text
                            if self._has_placeholders(original_cell_text):
                                logger.debug(f"[WORD-ERSETZUNG] Tabelle {table_idx}, Zeile {row_idx}, Zelle {cell_idx}: Gefundener Text: {original_cell_text[:200]}")
                                # Ersetze in allen Absätzen der Zelle
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        original_run_text = run.text
                                        run.text = self._replace_placeholders(run.text, extracted_info, process_id)
                                        if run.text != original_run_text:
                                            logger.info(f"[WORD-ERSETZUNG] Tabelle {table_idx}, Zelle {cell_idx}: Ersetzt '{original_run_text[:50]}' -> '{run.text[:50]}'")
                                            replacements_count += 1
            
            logger.info(f"[WORD-ERSETZUNG] Gesamtanzahl Ersetzungen: {replacements_count}")
            
            # Speichere neue Datei
            doc.save(str(output_path))
            logger.info(f"[WORD-ERSETZUNG] Word-Datei erstellt: {output_path} mit {replacements_count} Ersetzungen")
            return str(output_path)
            
        except ImportError:
            logger.error("python-docx ist nicht installiert")
            return None
        except Exception as e:
            logger.error(f"Fehler beim Erstellen der Word-Datei: {str(e)}")
            logger.exception(e)
            return None
    
    async def _create_xlsx_output(
        self,
        original_path: Path,
        output_path: Path,
        extracted_info: Dict[str, str],
        process_id: str
    ) -> Optional[str]:
        """Erstellt eine Excel-Datei mit ausgefüllten Platzhaltern"""
        try:
            from openpyxl import load_workbook
            
            logger.info(f"[EXCEL-ERSETZUNG] Lade Excel-Datei: {original_path}")
            logger.debug(f"[EXCEL-ERSETZUNG] Verfügbare Werte: {extracted_info}")
            
            # Lade Originalarbeitsmappe
            workbook = load_workbook(str(original_path))
            
            replacements_count = 0
            
            # Durchlaufe alle Arbeitsblätter
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                logger.debug(f"[EXCEL-ERSETZUNG] Verarbeite Arbeitsblatt: {sheet_name}")
                
                # Durchlaufe alle Zellen
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            cell_value = str(cell.value)
                            cell_coord = cell.coordinate
                            
                            # Prüfe, ob Platzhalter enthalten sind
                            if self._has_placeholders(cell_value):
                                logger.debug(f"[EXCEL-ERSETZUNG] Zelle {cell_coord}: Gefundener Wert mit Platzhaltern: {cell_value[:200]}")
                                # Ersetze Platzhalter
                                new_value = self._replace_placeholders(cell_value, extracted_info, process_id)
                                
                                if new_value != cell_value:
                                    logger.info(f"[EXCEL-ERSETZUNG] Zelle {cell_coord}: Ersetzt '{cell_value[:100]}' -> '{new_value[:100]}'")
                                    replacements_count += 1
                                
                                # Wenn es eine Formel war, versuche sie beizubehalten
                                if cell.data_type == 'f' or (isinstance(cell.value, str) and cell.value.startswith('=')):
                                    # Formel: Ersetze Platzhalter im Formel-String
                                    try:
                                        cell.value = new_value
                                    except Exception as e:
                                        logger.warning(f"[EXCEL-ERSETZUNG] Fehler beim Ersetzen in Formel {cell_coord}: {str(e)}")
                                        cell.value = new_value
                                else:
                                    # Normaler Wert: Ersetze direkt
                                    cell.value = new_value
            
            logger.info(f"[EXCEL-ERSETZUNG] Gesamtanzahl Ersetzungen: {replacements_count}")
            
            # Speichere neue Datei
            workbook.save(str(output_path))
            logger.info(f"[EXCEL-ERSETZUNG] Excel-Datei erstellt: {output_path} mit {replacements_count} Ersetzungen")
            return str(output_path)
            
        except ImportError:
            logger.error("openpyxl ist nicht installiert")
            return None
        except Exception as e:
            logger.error(f"Fehler beim Erstellen der Excel-Datei: {str(e)}")
            logger.exception(e)
            return None
    
    def _normalize_placeholder_name(self, name: str) -> str:
        """Normalisiert Platzhalter-Namen für Vergleich (entfernt Leerzeichen, normalisiert Groß-/Kleinschreibung)"""
        # Entferne führende/nachfolgende Leerzeichen
        normalized = name.strip()
        # Ersetze mehrere Leerzeichen durch ein einzelnes
        normalized = re.sub(r'\s+', ' ', normalized)
        return normalized
    
    def _find_matching_key(self, placeholder_name: str, extracted_info: Dict[str, str]) -> Optional[str]:
        """Findet einen passenden Key in extracted_info für den Platzhalter-Namen"""
        # Normalisiere Platzhalter-Name
        normalized_placeholder = self._normalize_placeholder_name(placeholder_name)
        
        # Direkter Match
        if normalized_placeholder in extracted_info:
            return normalized_placeholder
        
        # Fall-insensitive Match
        for key in extracted_info.keys():
            if self._normalize_placeholder_name(key).lower() == normalized_placeholder.lower():
                return key
        
        # Match ohne Leerzeichen
        placeholder_no_spaces = normalized_placeholder.replace(' ', '')
        for key in extracted_info.keys():
            key_no_spaces = self._normalize_placeholder_name(key).replace(' ', '')
            if key_no_spaces.lower() == placeholder_no_spaces.lower():
                return key
        
        return None
    
    def _replace_placeholders(self, text: str, extracted_info: Dict[str, str], process_id: str = None) -> str:
        """Ersetzt Platzhalter im Format {{feldname}} durch Werte aus extracted_info"""
        if not text:
            return text
        
        log_prefix = f"[PLATZHALTER-ERSETZUNG-{process_id}]" if process_id else "[PLATZHALTER-ERSETZUNG]"
        
        result = text
        # Finde alle Platzhalter im Format {{feldname}} - unterstützt auch Leerzeichen
        # Pattern: {{...}} wobei ... beliebige Zeichen außer } sein können
        pattern = r'\{\{([^}]+)\}\}'
        
        # Finde alle Platzhalter im Text
        found_placeholders = re.findall(pattern, text)
        if found_placeholders:
            logger.debug(f"{log_prefix} Gefundene Platzhalter im Text '{text[:200]}': {found_placeholders}")
            logger.debug(f"{log_prefix} Verfügbare Werte in extracted_info: {list(extracted_info.keys())}")
        
        def replace_match(match):
            placeholder_name = match.group(1).strip()  # Entferne Leerzeichen am Anfang/Ende
            logger.debug(f"{log_prefix} Suche nach Platzhalter: '{placeholder_name}'")
            
            # Versuche passenden Key zu finden
            matching_key = self._find_matching_key(placeholder_name, extracted_info)
            
            if matching_key:
                replacement_value = str(extracted_info[matching_key])
                logger.info(f"{log_prefix} Ersetze '{{{{{placeholder_name}}}}}' (gematcht mit Key '{matching_key}') durch '{replacement_value[:100]}'")
                return replacement_value
            else:
                logger.warning(f"{log_prefix} Platzhalter '{{{{{placeholder_name}}}}}' nicht in extracted_info gefunden!")
                logger.debug(f"{log_prefix} Verfügbare Keys: {list(extracted_info.keys())}")
                logger.debug(f"{log_prefix} Normalisierter Platzhalter: '{self._normalize_placeholder_name(placeholder_name)}'")
                # Wenn Platzhalter nicht gefunden, lasse Original stehen
                return match.group(0)
        
        result = re.sub(pattern, replace_match, result)
        
        if result != text:
            logger.debug(f"{log_prefix} Text geändert: '{text[:200]}' -> '{result[:200]}'")
        else:
            logger.debug(f"{log_prefix} Keine Ersetzung durchgeführt für Text: '{text[:200]}'")
        
        return result
    
    def _has_placeholders(self, text: str) -> bool:
        """Prüft, ob Text Platzhalter im Format {{feldname}} enthält (unterstützt auch Leerzeichen)"""
        if not text:
            return False
        # Pattern unterstützt jetzt auch Leerzeichen und andere Zeichen
        pattern = r'\{\{([^}]+)\}\}'
        return bool(re.search(pattern, text))
